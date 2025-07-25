# converter.py
# (Este es el núcleo de tu script original, sin la parte de tkinter)
import os
import re
# Imports para DOCX
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

class MarkdownConverter:
    def __init__(self):
        self.debug = False

    def set_debug(self, debug):
        self.debug = debug

    def log(self, message):
        if self.debug:
            print(f"[DEBUG] {message}")

    def parse_markdown_elements(self, content):
        """
        Parsea el contenido Markdown y devuelve una lista de elementos estructurados.
        """
        lines = content.splitlines()
        elements = []
        in_code_block = False
        code_block_content = []
        in_table = False
        table_rows = []
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            # Bloques de código
            if stripped.startswith("```"):
                if in_code_block:
                    # Fin del bloque de código
                    elements.append({
                        'type': 'code_block',
                        'content': '\n'.join(code_block_content), # Corregido salto de línea
                        'language': ''
                    })
                    code_block_content = []
                    in_code_block = False
                else:
                    # Inicio del bloque de código
                    in_code_block = True
                i += 1
                continue
            if in_code_block:
                code_block_content.append(line)
                i += 1
                continue
            # Detectar tablas (líneas con |)
            if '|' in stripped and not in_table:
                # Comprobar si la siguiente línea es separador de tabla
                if i + 1 < len(lines) and re.match(r'^\s*\|?[-:|\s]+\|?\s*$', lines[i + 1].strip()):
                    in_table = True
                    table_rows = []
                    # Agregar encabezado
                    headers = [cell.strip() for cell in stripped.split('|') if cell.strip()]
                    table_rows.append(headers)
                    i += 2  # Saltar línea separadora
                    continue
            if in_table and '|' in stripped:
                # Agregar fila de tabla
                cells = [cell.strip() for cell in stripped.split('|') if cell.strip()]
                table_rows.append(cells)
                i += 1
                continue
            elif in_table:
                # Fin de tabla
                elements.append({
                    'type': 'table',
                    'rows': table_rows
                })
                in_table = False
                table_rows = []
                # No incrementar i, procesar esta línea normalmente
            # Encabezados
            header_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2)
                elements.append({
                    'type': 'heading',
                    'level': level,
                    'content': text
                })
                i += 1
                continue
            # Separadores
            if re.match(r'^-{3,}', stripped):
                elements.append({'type': 'separator'})
                i += 1
                continue
            # Citas
            if stripped.startswith(">"):
                quote_text = stripped[1:].strip()
                elements.append({
                    'type': 'quote',
                    'content': quote_text
                })
                i += 1
                continue
            # Listas
            list_match = re.match(r'^([-+*]|\d+\.)\s*(.*)', stripped)
            if list_match:
                marker = list_match.group(1)
                content = list_match.group(2)
                list_type = 'ordered' if re.match(r'^\d+\.', marker) else 'unordered'
                elements.append({
                    'type': 'list_item',
                    'list_type': list_type,
                    'content': content
                })
                i += 1
                continue
            # Líneas vacías
            if not stripped:
                elements.append({'type': 'empty_line'})
                i += 1
                continue
            # Texto normal
            elements.append({
                'type': 'paragraph',
                'content': stripped
            })
            i += 1
        # Si terminamos en una tabla, agregarla
        if in_table and table_rows:
            elements.append({
                'type': 'table',
                'rows': table_rows
            })
        return elements

    def parse_inline_styles(self, text):
        """
        Parsea estilos en línea: **negrita**, *cursiva*, `código`
        Retorna lista de tuplas (texto, estilo)
        """
        pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`)'
        parts = re.split(pattern, text)
        result = []
        for part in parts:
            if not part:
                continue
            elif part.startswith('**') and part.endswith('**'):
                result.append((part[2:-2], 'bold'))
            elif part.startswith('*') and part.endswith('*'):
                result.append((part[1:-1], 'italic'))
            elif part.startswith('`') and part.endswith('`'):
                result.append((part[1:-1], 'code'))
            else:
                result.append((part, 'normal'))
        return result

    def convert_to_docx(self, elements, output_path, template_path=None):
        """Convierte elementos a formato DOCX, opcionalmente usando una plantilla."""
        # Usar plantilla si se proporciona, de lo contrario crear documento nuevo
        if template_path and os.path.exists(template_path):
            try:
                doc = Document(template_path)
                self.log(f"Usando plantilla: {template_path}")
            except Exception as e:
                self.log(f"Advertencia: No se pudo cargar la plantilla {template_path}: {e}. Creando documento nuevo.")
                doc = Document()
        else:
            doc = Document()
            self.log("Creando nuevo documento DOCX")

        # --- Manejo robusto de encabezados ---
        # Obtener lista de estilos disponibles
        styles = doc.styles
        available_styles = [s.name for s in styles]

        # --- Crear estilos personalizados si no existen ---
        if 'Code' not in available_styles:
            try:
                code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
                code_font = code_style.font
                code_font.name = 'Courier New'
                code_font.size = Pt(10)
                self.log("Estilo 'Code' creado.")
            except Exception as e:
                self.log(f"Advertencia al crear estilo 'Code': {e}")
        if 'Code Char' not in available_styles:
            try:
                code_char_style = styles.add_style('Code Char', WD_STYLE_TYPE.CHARACTER)
                code_char_font = code_char_style.font
                code_char_font.name = 'Courier New'
                code_char_font.size = Pt(10)
                self.log("Estilo 'Code Char' creado.")
            except Exception as e:
                 self.log(f"Advertencia al crear estilo 'Code Char': {e}")

        # --- Insertar contenido convertido al final ---
        for element in elements:
            if element['type'] == 'heading':
                level = min(max(element['level'], 1), 9) # Asegurar nivel válido (1-9)
                # Determinar el nombre del estilo de encabezado
                heading_style_name = f"Heading {level}"
                # Verificar si el estilo existe en la plantilla/cargado
                if heading_style_name in available_styles:
                    # Usar el estilo definido en la plantilla/documento
                    doc.add_heading(element['content'], level)
                else:
                    # Fallback: Crear párrafo con estilo Normal y aplicar formato básico
                    self.log(f"Advertencia: Estilo '{heading_style_name}' no encontrado. Usando estilo Normal con formato.")
                    p = doc.add_paragraph(element['content'])
                    # Aplicar formato básico: negrita y tamaño de fuente proporcional
                    if p.runs: # Verificar si hay runs
                        run = p.runs[0]
                        # Asignar tamaño de fuente basado en el nivel del encabezado
                        font_size_map = {1: 20, 2: 18, 3: 16, 4: 14, 5: 12, 6: 11}
                        run.font.size = Pt(font_size_map.get(level, 12))
                        # Aplicar negrita
                        run.bold = True
            elif element['type'] == 'paragraph':
                p = doc.add_paragraph()
                fragments = self.parse_inline_styles(element['content'])
                for text, style in fragments:
                    run = p.add_run(text)
                    if style == 'bold':
                        run.bold = True
                    elif style == 'italic':
                        run.italic = True
                    elif style == 'code':
                        # Aplicar estilo de carácter para código inline
                        if 'Code Char' in available_styles:
                             run.style = 'Code Char'
                        else:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(10)
            elif element['type'] == 'table':
                # Crear tabla simple en DOCX
                if element['rows']: # Verificar que haya filas
                    # Asegurarse de que todas las filas tengan el mismo número de columnas
                    max_cols = max(len(row) for row in element['rows']) if element['rows'] else 0
                    table = doc.add_table(rows=len(element['rows']), cols=max_cols)
                    table.style = 'Table Grid'
                    for row_index, row_data in enumerate(element['rows']):
                        for col_index, cell_data in enumerate(row_data):
                            # Asegurarse de no exceder el número de columnas de la tabla
                            if col_index < max_cols:
                                cell = table.cell(row_index, col_index)
                                cell_paragraph = cell.paragraphs[0]
                                cell_paragraph.text = cell_data
                                # Si es encabezado (primera fila), aplicar negrita
                                if row_index == 0:
                                    for run in cell_paragraph.runs:
                                        run.bold = True
            elif element['type'] == 'code_block':
                p = doc.add_paragraph(element['content'])
                # Aplicar estilo de párrafo para bloques de código
                if 'Code' in available_styles:
                    p.style = 'Code'
                else:
                     for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
            elif element['type'] == 'quote':
                p = doc.add_paragraph(element['content'])
                # Usar estilo Quote si existe, sino usar Normal
                if 'Quote' in available_styles:
                    p.style = 'Quote'
                # Si no existe, se deja con estilo Normal
            elif element['type'] == 'list_item':
                if element['list_type'] == 'ordered':
                    p = doc.add_paragraph(element['content'], style='List Number')
                else:
                    p = doc.add_paragraph(element['content'], style='List Bullet')
            elif element['type'] == 'separator':
                p = doc.add_paragraph("―" * 30)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif element['type'] == 'empty_line':
                doc.add_paragraph()
        doc.save(output_path)
        self.log(f"Archivo DOCX guardado: {output_path}")

# Si este archivo se ejecuta directamente, puedes hacer pruebas aquí
if __name__ == "__main__":
    pass # Las pruebas se harían en app.py o en un script aparte
